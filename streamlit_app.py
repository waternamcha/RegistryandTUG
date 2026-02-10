import streamlit as st
import time
from datetime import datetime
import io

# 📦 ต้องติดตั้ง: pip install python-docx
from docx import Document 
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# =========================================================
# ⚙️ CONFIG & CSS
# =========================================================
st.set_page_config(page_title="Prosthesis Registry & TUG", page_icon="🦿", layout="wide")

st.markdown("""
    <style>
    /* ปรับตัวเลขนาฬิกาให้ใหญ่และอ่านง่าย */
    div[data-testid="stMetricValue"] {
        font-size: 80px !important;
        font-family: 'Courier New', monospace;
        color: #1F618D;
        font-weight: bold;
    }
    /* กรอบ TUG */
    .tug-box {
        border: 2px solid #ddd;
        padding: 20px;
        border-radius: 10px;
        background-color: #f9f9f9;
        text-align: center;
    }
    .stButton button {
        height: 3em;
        font-weight: bold;
    }
    </style>
""", unsafe_allow_html=True)

# =========================================================
# 📦 SESSION STATE (เตรียมตัวแปรให้ครบ 33 ข้อ)
# =========================================================
defaults = {
    # TUG System
    'is_running': False, 
    'start_time': 0.0, 
    'elapsed_time': 0.0,
    
    # 1. General
    'hn': '', 'fname': '', 'birth_year': 2520, 'gender': 'ชาย', 
    'weight': 0.0, 'height': 0.0, 'nationality': 'ไทย', 'nat_ot': '',
    'country': 'Thailand', 'cnt_ot': '', 'province': 'กรุงเทพมหานคร',
    
    # 2. Medical
    'comorbidities': [], 'comorb_ot': '',
    'cause': 'อุบัติเหตุ', 'cause_ot': '',
    'amp_year': 2566, 'side': 'ขวา',
    'amp_level': 'Transtibial', 'level_ot': '',
    'stump_len': 'ปานกลาง', 'stump_shape': 'Cylindrical', 'shape_ot': '',
    'surgery': 'ไม่ใช่', 'surg_details': [], 'surg_ot': '',
    'k_level': 'K3',
    
    # 3. Rehab
    'personnel': [], 'pers_ot': '',
    'rehab_status': 'ไม่เคย', 'activities': [], 'act_ot': '',
    
    # 4. Prosthesis
    'service': [], 'serv_ot': '',
    'date_cast': datetime.now().date(), 'date_deliv': datetime.now().date(),
    'socket': 'PTB', 'sock_ot': '',
    'liner': [], 'liner_ot': '',
    'suspension': [], 'susp_ot': '',
    'foot': [], 'foot_ot': '',
    'knee': [], 'knee_ot': '', # Item 27 (เฉพาะเหนือเข่า)
    
    # 5. Social & Function (Items 28-33)
    'assist': 'ไม่ใช้', 'asst_ot': '',
    'stand_hours': '1-3 ชม.', 'walk_hours': '1-3 ชม.',
    'fall_hist': 'ไม่มี', 'fall_freq': '1-2 ครั้ง', 'fall_inj': False,
    # Q31-32
    'q31_1': 'ไม่มีปัญหา (0-4%)', 'q31_2': 'ไม่มีปัญหา (0-4%)', 
    'q32_1': 'ไม่มีปัญหา (0-4%)', 'q32_2': 'ไม่มีปัญหา (0-4%)',
    # Q33 Support
    'supp_family': 'ใช่', 
    'supp_org': 'ไม่ใช่', 'supp_sources': [], 'supp_ot': '',
    
    # TUG Results
    't1': 0.0, 't2': 0.0, 't3': 0.0, 'tug_avg': 0.0
}

for key, val in defaults.items():
    if key not in st.session_state:
        st.session_state[key] = val

# ฟังก์ชันช่วยจัดการ Text ที่มี "Other"
def get_val(key, other_key=None):
    val = st.session_state[key]
    if isinstance(val, list):
        res = ", ".join(val)
        if "Other" in val and other_key and st.session_state[other_key]:
            res += f" ({st.session_state[other_key]})"
        return res if res else "-"
    else:
        if val == "Other" and other_key and st.session_state[other_key]:
            return f"{st.session_state[other_key]} (Other)"
        return str(val)

# =========================================================
# 📄 WORD GENERATOR (Fixed Complete Data)
# =========================================================
def generate_word_report():
    doc = Document()
    
    # Config Font (ใช้ Angsana New เพื่อภาษาไทยที่ถูกต้อง)
    style = doc.styles['Normal']
    style.font.name = 'Angsana New'
    style.font.size = Pt(14)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Angsana New')

    # Title
    head = doc.add_heading('แบบสำรวจประวัติผู้ใช้ขาเทียม (Registry Report)', 0)
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"พิมพ์เมื่อ: {datetime.now().strftime('%d/%m/%Y %H:%M')}\n").alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- 1. ข้อมูลทั่วไป ---
    doc.add_heading('1. ข้อมูลทั่วไป (General)', level=1)
    t1 = doc.add_table(rows=3, cols=2)
    t1.style = 'Table Grid'
    r = t1.rows
    r[0].cells[0].text = f"ชื่อ-สกุล: {st.session_state.fname}"
    r[0].cells[1].text = f"HN: {st.session_state.hn}"
    r[1].cells[0].text = f"เพศ: {st.session_state.gender} | เกิด: {st.session_state.birth_year} (อายุ {datetime.now().year + 543 - st.session_state.birth_year} ปี)"
    r[1].cells[1].text = f"รูปร่าง: {st.session_state.weight} กก. / {st.session_state.height} ซม."
    r[2].cells[0].text = f"ที่อยู่: {st.session_state.province}, {get_val('country', 'cnt_ot')}"
    r[2].cells[1].text = f"สัญชาติ: {get_val('nationality', 'nat_ot')}"

    # --- 2. การแพทย์ ---
    doc.add_heading('2. ข้อมูลทางการแพทย์ (Medical)', level=1)
    t2 = doc.add_table(rows=5, cols=2)
    t2.style = 'Table Grid'
    r = t2.rows
    r[0].cells[0].text = f"โรคประจำตัว: {get_val('comorbidities', 'comorb_ot')}"
    r[0].cells[1].text = f"สาเหตุการตัด: {get_val('cause', 'cause_ot')}"
    r[1].cells[0].text = f"ระดับการตัด: {get_val('amp_level', 'level_ot')} ({st.session_state.side})"
    r[1].cells[1].text = f"ปีที่ตัด: {st.session_state.amp_year} (K-Level: {st.session_state.k_level})"
    r[2].cells[0].text = f"ตอขา: {st.session_state.stump_len}, {get_val('stump_shape', 'shape_ot')}"
    r[2].cells[1].text = f"ผ่าตัดเพิ่มเติม: {st.session_state.surgery} {get_val('surg_details', 'surg_ot')}"
    
    # --- 3. ฟื้นฟู ---
    doc.add_heading('3. การฟื้นฟู (Rehab)', level=1)
    r[3].cells[0].text = f"บุคลากร: {get_val('personnel', 'pers_ot')}"
    r[3].cells[1].text = f"สถานะฟื้นฟู: {st.session_state.rehab_status} ({get_val('activities', 'act_ot')})"

    # --- 4. กายอุปกรณ์ ---
    doc.add_heading('4. ข้อมูลกายอุปกรณ์ (Prosthesis)', level=1)
    p = doc.add_paragraph()
    p.add_run(f"บริการครั้งนี้: {get_val('service', 'serv_ot')}\n").bold = True
    p.add_run(f"วันที่หล่อ: {st.session_state.date_cast} | วันที่รับ: {st.session_state.date_deliv}\n")
    p.add_run(f"Socket: {get_val('socket', 'sock_ot')} | Liner: {get_val('liner', 'liner_ot')}\n")
    p.add_run(f"Suspension: {get_val('suspension', 'susp_ot')} | Foot: {get_val('foot', 'foot_ot')}\n")
    
    # Logic: โชว์เข่าเฉพาะเมื่อตัดขาเหนือเข่า
    if st.session_state.amp_level in ["Transfemoral", "Knee Disarticulation", "Other"]:
        knee_txt = get_val('knee', 'knee_ot')
        p.add_run(f"Knee (ข้อเข่า): {knee_txt}").bold = True
    else:
        p.add_run("Knee: - (ระดับต่ำกว่าเข่า หรือไม่ได้ระบุ)")

    # --- 5. สังคม & การใช้งาน ---
    doc.add_heading('5. สังคมและการใช้งาน (Social & Function)', level=1)
    t3 = doc.add_table(rows=2, cols=2)
    t3.style = 'Table Grid'
    r = t3.rows
    r[0].cells[0].text = f"อุปกรณ์ช่วยเดิน: {get_val('assist', 'asst_ot')}"
    r[0].cells[1].text = f"การใช้งาน (ยืน/เดิน): {st.session_state.stand_hours} / {st.session_state.walk_hours}"
    
    fall_txt = "ไม่มี"
    if st.session_state.fall_hist == "มี":
        inj = "บาดเจ็บ" if st.session_state.fall_inj else "ไม่เจ็บ"
        fall_txt = f"มี ({st.session_state.fall_freq}) - {inj}"
    r[1].cells[0].text = f"ประวัติล้ม (6ด.): {fall_txt}"
    r[1].cells[1].text = f"ดูแลจากครอบครัว: {st.session_state.supp_family}"

    # ส่วนที่เคยหายไป (Q31-33)
    doc.add_paragraph("\nการประเมินตนเอง (Items 31-32):").bold = True
    doc.add_paragraph(f"• สังคม (ตนเอง): {st.session_state.q31_1}")
    doc.add_paragraph(f"• สังคม (เทียบคนอื่น): {st.session_state.q31_2}")
    doc.add_paragraph(f"• งาน (ตนเอง): {st.session_state.q32_1}")
    doc.add_paragraph(f"• งาน (เทียบคนอื่น): {st.session_state.q32_2}")
    
    supp_txt = "ไม่มี"
    if st.session_state.supp_org == "ใช่":
        supp_txt = get_val('supp_sources', 'supp_ot')
    doc.add_paragraph(f"การสนับสนุนจากหน่วยงาน (Item 33): {supp_txt}")

    # --- 6. TUG Results ---
    doc.add_heading('6. ผลทดสอบ TUG Test', level=1)
    res_msg = "NORMAL" if st.session_state.tug_avg < 13.5 else "HIGH RISK"
    
    p_tug = doc.add_paragraph()
    p_tug.add_run(f"Trial 1: {st.session_state.t1} s\n")
    p_tug.add_run(f"Trial 2: {st.session_state.t2} s\n")
    p_tug.add_run(f"Trial 3: {st.session_state.t3} s\n")
    run = p_tug.add_run(f"\nAverage: {st.session_state.tug_avg:.2f} seconds ({res_msg})")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 128, 0) if st.session_state.tug_avg < 13.5 else RGBColor(200, 0, 0)

    bio = io.BytesIO()
    doc.save(bio)
    return bio

# =========================================================
# 📱 APP UI & LOGIC
# =========================================================

st.sidebar.title("🦿 เมนูหลัก")
st.sidebar.info("แนะนำ: กรอกข้อมูลให้ครบก่อนกดดาวน์โหลด")

if st.sidebar.button("📄 สร้างไฟล์ Word (.docx)"):
    if not st.session_state.hn:
        st.sidebar.error("⚠️ กรุณากรอก HN ก่อน")
    else:
        file_buffer = generate_word_report()
        st.sidebar.download_button(
            label="⬇️ ดาวน์โหลดไฟล์ Word",
            data=file_buffer,
            file_name=f"Report_{st.session_state.hn}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

tab1, tab2 = st.tabs(["📝 กรอกข้อมูล (Registry)", "⏱️ จับเวลา (TUG Test)"])

with tab1:
    st.header("แบบสำรวจข้อมูล (Items 1-33)")
    
    with st.expander("1. ข้อมูลทั่วไป", expanded=True):
        c1, c2, c3 = st.columns(3)
        st.session_state.hn = c1.text_input("HN", key="i_hn")
        st.session_state.fname = c1.text_input("ชื่อ-นามสกุล", key="i_fname")
        st.session_state.birth_year = c1.selectbox("ปีเกิด", list(range(2567, 2467, -1)), key="i_byear")
        st.session_state.gender = c2.selectbox("เพศ", ["ชาย", "หญิง"], key="i_gen")
        st.session_state.province = c2.selectbox("จังหวัด", ["กรุงเทพมหานคร", "เชียงใหม่", "ขอนแก่น", "อื่นๆ"], key="i_prov")
        st.session_state.nationality = c2.selectbox("สัญชาติ", ["ไทย", "Other"], key="i_nat")
        if st.session_state.nationality=="Other": st.session_state.nat_ot = c2.text_input("ระบุสัญชาติ", key="i_not")
        st.session_state.weight = c3.number_input("น้ำหนัก (กก.)", 0.0, key="i_wt")
        st.session_state.height = c3.number_input("ส่วนสูง (ซม.)", 0.0, key="i_ht")
        st.session_state.country = c3.selectbox("ประเทศ", ["Thailand", "Other"], key="i_cnt")
        if st.session_state.country=="Other": st.session_state.cnt_ot = c3.text_input("ระบุประเทศ", key="i_cot")

    with st.expander("2. ข้อมูลการแพทย์"):
        c1, c2 = st.columns(2)
        st.session_state.comorbidities = c1.multiselect("โรคประจำตัว", ["เบาหวาน", "ความดัน", "หัวใจ", "ไม่มี", "Other"], key="i_com")
        if "Other" in st.session_state.comorbidities: st.session_state.comorb_ot = c1.text_input("ระบุโรค", key="i_com_ot")
        st.session_state.cause = c1.selectbox("สาเหตุตัดขา", ["อุบัติเหตุ", "เบาหวาน", "มะเร็ง", "Other"], key="i_cau")
        if st.session_state.cause=="Other": st.session_state.cause_ot = c1.text_input("ระบุสาเหตุ", key="i_cau_ot")
        st.session_state.amp_level = c2.selectbox("ระดับการตัด", ["Transtibial", "Transfemoral", "Knee Disarticulation", "Other"], key="i_lvl")
        if st.session_state.amp_level=="Other": st.session_state.level_ot = c2.text_input("ระบุระดับ", key="i_lvl_ot")
        st.session_state.side = c2.radio("ข้าง", ["ซ้าย", "ขวา", "สองข้าง"], horizontal=True, key="i_side")
        st.session_state.amp_year = c2.number_input("ปีที่ตัด (พ.ศ.)", 2490, 2600, key="i_ayr")
        st.session_state.k_level = c2.selectbox("K-Level", ["K0", "K1", "K2", "K3", "K4"], key="i_k")
        st.session_state.stump_len = c1.selectbox("ความยาวตอขา", ["สั้น", "ปานกลาง", "ยาว"], key="i_slen")
        st.session_state.stump_shape = c1.selectbox("รูปทรงตอขา", ["Conical", "Cylindrical", "Bulbous", "Other"], key="i_shp")
        if st.session_state.stump_shape=="Other": st.session_state.shape_ot = c1.text_input("ระบุทรง", key="i_shp_ot")
        st.session_state.surgery = c2.radio("ผ่าตัดเพิ่มเติม", ["ไม่ใช่", "ใช่"], key="i_surg")
        if st.session_state.surgery=="ใช่":
            st.session_state.surg_details = c2.multiselect("ระบุ", ["ตัดกระดูก", "ตัดผิวหนัง", "Other"], key="i_surg_d")
            if "Other" in st.session_state.surg_details: st.session_state.surg_ot = c2.text_input("ระบุผ่าตัด", key="i_surg_ot")

    with st.expander("3-4. การฟื้นฟู & กายอุปกรณ์"):
        st.session_state.personnel = st.multiselect("บุคลากร", ["นักกายอุปกรณ์", "นักกายภาพ", "แพทย์", "Other"], key="i_per")
        if "Other" in st.session_state.personnel: st.session_state.pers_ot = st.text_input("ระบุบุคลากร", key="i_per_ot")
        st.session_state.rehab_status = st.radio("ประวัติฟื้นฟู", ["ไม่เคย", "เคย"], horizontal=True, key="i_reh")
        if st.session_state.rehab_status=="เคย":
            st.session_state.activities = st.multiselect("กิจกรรม", ["ถุงลดบวม", "พันผ้ายืด", "ฝึกเดิน", "Other"], key="i_act")
            if "Other" in st.session_state.activities: st.session_state.act_ot = st.text_input("ระบุกิจกรรม", key="i_act_ot")
        st.markdown("---")
        st.session_state.service = st.multiselect("บริการ", ["ทำใหม่", "เปลี่ยนเบ้า", "ซ่อม", "Other"], key="i_srv")
        if "Other" in st.session_state.service: st.session_state.serv_ot = st.text_input("ระบุบริการ", key="i_srv_ot")
        c1, c2 = st.columns(2)
        st.session_state.date_cast = c1.date_input("วันหล่อแบบ", key="i_dc")
        st.session_state.date_deliv = c2.date_input("วันรับขา", key="i_dd")
        st.session_state.socket = c1.selectbox("Socket", ["PTB", "TSB", "Ischial", "Other"], key="i_sck")
        if st.session_state.socket=="Other": st.session_state.sock_ot = c1.text_input("ระบุ Socket", key="i_sck_ot")
        st.session_state.liner = c1.multiselect("Liner", ["No liner", "Foam", "Silicone", "Other"], key="i_lin")
        if "Other" in st.session_state.liner: st.session_state.liner_ot = c1.text_input("ระบุ Liner", key="i_lin_ot")
        st.session_state.suspension = c2.multiselect("Suspension", ["Suction", "Pin lock", "Belt", "Other"], key="i_sus")
        if "Other" in st.session_state.suspension: st.session_state.susp_ot = c2.text_input("ระบุ Susp", key="i_sus_ot")
        st.session_state.foot = c2.multiselect("Foot", ["SACH", "Single axis", "Dynamic", "Other"], key="i_ft")
        if "Other" in st.session_state.foot: st.session_state.foot_ot = c2.text_input("ระบุ Foot", key="i_ft_ot")

        # Knee Section: โชว์เฉพาะถ้าตัดเหนือเข่า
        if st.session_state.amp_level in ["Transfemoral", "Knee Disarticulation", "Other"]:
            st.info("🦵 ส่วนนี้สำหรับระดับเหนือเข่า (Transfemoral) ขึ้นไป")
            st.session_state.knee = st.multiselect("เลือก Knee", ["Single axis", "Polycentric", "Hydraulic", "Other"], key="i_kn")
            if "Other" in st.session_state.knee: st.session_state.knee_ot = st.text_input("ระบุ Knee", key="i_kn_ot")

    with st.expander("5. สังคม & การใช้งาน", expanded=True):
        c1, c2 = st.columns(2)
        st.session_state.assist = c1.selectbox("อุปกรณ์ช่วยเดิน", ["ไม่ใช้", "ไม้เท้า", "Walker", "Other"], key="i_ast")
        if st.session_state.assist=="Other": st.session_state.asst_ot = c1.text_input("ระบุอุปกรณ์", key="i_ast_ot")
        st.session_state.stand_hours = c1.selectbox("ยืน/วัน", ["< 1 ชม.", "1-3 ชม.", "> 3 ชม."], key="i_std")
        st.session_state.walk_hours = c2.selectbox("เดิน/วัน", ["< 1 ชม.", "1-3 ชม.", "> 3 ชม."], key="i_wlk")
        
        st.markdown("---")
        st.session_state.fall_hist = st.radio("ประวัติล้ม (6ด.)", ["ไม่มี", "มี"], horizontal=True, key="i_fall")
        if st.session_state.fall_hist=="มี":
            c1, c2 = st.columns(2)
            st.session_state.fall_freq = c1.selectbox("ความถี่การล้ม", ["1-2 ครั้ง", "> 2 ครั้ง"], key="i_ffrq")
            st.session_state.fall_inj = c2.checkbox("มีการบาดเจ็บ", key="i_finj")
        
        st.markdown("---")
        st.write("##### 31-32. การมีส่วนร่วมในสังคม & การทำงาน")
        p_lvl = ["ไม่มีปัญหา (0-4%)", "มีปัญหาเล็กน้อย (5-24%)", "มีปัญหาปานกลาง (25-49%)", "มีปัญหามาก (50-95%)", "มีปัญหามากที่สุด (96-100%)"]
        c1, c2 = st.columns(2)
        st.session_state.q31_1 = c1.selectbox("31.1 สังคม (ตนเอง)", p_lvl, key="i_q311")
        st.session_state.q31_2 = c2.selectbox("31.2 สังคม (เทียบคนอื่น)", p_lvl, key="i_q312")
        st.session_state.q32_1 = c1.selectbox("32.1 งาน (ตนเอง)", p_lvl, key="i_q321")
        st.session_state.q32_2 = c2.selectbox("32.2 งาน (เทียบคนอื่น)", p_lvl, key="i_q322")
        
        st.markdown("---")
        st.write("##### 33. การสนับสนุน")
        st.session_state.supp_family = st.radio("33.1 การดูแลจากครอบครัว", ["ใช่", "ไม่ใช่"], horizontal=True, key="i_sfam")
        st.session_state.supp_org = st.radio("33.2 สนับสนุนจากหน่วยงาน", ["ใช่", "ไม่ใช่"], horizontal=True, key="i_sorg")
        if st.session_state.supp_org=="ใช่":
            st.session_state.supp_sources = st.multiselect("ระบุหน่วยงาน", ["รัฐ", "เอกชน", "Other"], key="i_ssrc")
            if "Other" in st.session_state.supp_sources: st.session_state.supp_ot = st.text_input("ระบุหน่วยงานอื่น", key="i_ssrc_ot")

with tab2:
    st.header("⏱️ Timed Up and Go (TUG)")
    
    # --- LOGIC นาฬิกาใหม่ (ไม่ค้าง) ---
    col_clock, col_btns = st.columns([2, 1])
    
    with col_clock:
        clock_container = st.empty()
        # แสดงเวลา
        if st.session_state.is_running:
            # คำนวณเวลาสดจาก start_time
            now = time.time()
            elapsed = now - st.session_state.start_time
            clock_container.metric(label="Time (Seconds)", value=f"{elapsed:.2f} s")
            # Rerun เพื่อให้เวลาเดิน (แต่ปุ่ม Stop จะทำงานได้เพราะอยู่นอก Container นี้)
            time.sleep(0.1) 
            st.rerun()
        else:
            clock_container.metric(label="Time (Seconds)", value=f"{st.session_state.elapsed_time:.2f} s")

    with col_btns:
        # ปุ่มควบคุมแยกออกมาเพื่อให้กดติดง่าย
        if st.button("▶️ START", type="primary", use_container_width=True, disabled=st.session_state.is_running):
            st.session_state.is_running = True
            st.session_state.start_time = time.time()
            st.rerun()

        if st.button("⏹️ STOP", type="secondary", use_container_width=True, disabled=not st.session_state.is_running):
            st.session_state.is_running = False
            st.session_state.elapsed_time = time.time() - st.session_state.start_time
            st.rerun()

        if st.button("🔄 RESET", use_container_width=True):
            st.session_state.is_running = False
            st.session_state.elapsed_time = 0.0
            st.rerun()

    st.markdown("---")
    st.write("บันทึกผลการทดสอบ:")
    c1, c2, c3 = st.columns(3)
    st.session_state.t1 = c1.number_input("Trial 1", 0.0, key="v_t1")
    st.session_state.t2 = c2.number_input("Trial 2", 0.0, key="v_t2")
    st.session_state.t3 = c3.number_input("Trial 3", 0.0, key="v_t3")

    valid_times = [t for t in [st.session_state.t1, st.session_state.t2, st.session_state.t3] if t > 0]
    if valid_times:
        st.session_state.tug_avg = sum(valid_times) / len(valid_times)
        status = "High Fall Risk (เสี่ยงล้มสูง)" if st.session_state.tug_avg >= 13.5 else "Normal Mobility (ปกติ)"
        color = "#C0392B" if st.session_state.tug_avg >= 13.5 else "#28B463"
        st.markdown(f"""
        <div class="tug-box" style="border-color: {color}; color: {color};">
            <h3>Average: {st.session_state.tug_avg:.2f} sec</h3>
            <h1>{status}</h1>
        </div>
        """, unsafe_allow_html=True)